#!/usr/bin/env python3
"""
Create a full Kaifa manuscript DOCX from the supplied source DOCX.

This script preserves the original Word-manuscript layout, figures, and
reference list while updating the core text and SEM tables to match the
working SEM evidence base archived in this repository.
"""

from __future__ import annotations

import argparse
from copy import deepcopy
from pathlib import Path
import sys

from docx import Document
import pandas as pd

REPO_ROOT = Path(__file__).resolve().parents[2]
SRC_ROOT = REPO_ROOT / "src"
KAIFA_DIAGNOSTICS_DIR = (
    REPO_ROOT / "data_work" / "diagnostics" / "kaifa_recovered_analysis"
)


def clear_paragraph(paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            p.remove(child)


def replace_paragraph_text(paragraph, text: str) -> None:
    first_run = paragraph.runs[0] if paragraph.runs else None
    attrs = {
        "bold": first_run.bold if first_run else None,
        "italic": first_run.italic if first_run else None,
        "underline": first_run.underline if first_run else None,
    }
    font = None
    if first_run is not None:
        font = {
            "name": first_run.font.name,
            "size": deepcopy(first_run.font.size),
        }

    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.bold = attrs["bold"]
    run.italic = attrs["italic"]
    run.underline = attrs["underline"]
    if font is not None:
        run.font.name = font["name"]
        run.font.size = font["size"]


def update_cell(cell, text: str) -> None:
    if not cell.paragraphs:
        cell.text = text
        return
    replace_paragraph_text(cell.paragraphs[0], text)


def normalize_empty_paragraph(paragraph) -> None:
    """Strip runs and paragraph properties from empty spacer paragraphs."""
    p = paragraph._element
    for child in list(p):
        p.remove(child)
    paragraph.add_run("")


PARAGRAPH_REPLACEMENTS = {
    5: (
        "This study analyzes 573 cross-sectional administering-jurisdiction profiles "
        "(30 state agencies and 543 county-linked local jurisdictions) constructed "
        "from HUD DRGR/QPR disaster-recovery records spanning disasters from 2003 "
        "through 2023. Linked QCEW staffing and payroll proxies, population measures, "
        "and CDC/ATSDR social vulnerability indicators are used in one-factor and "
        "two-factor structural equation models that summarize administrative resources, "
        "workload manageability, recovery performance, and recovery timeliness. The "
        "most consistent pattern in the preferred two-factor model is that more "
        "manageable workload conditions are associated with faster Recovery Timeliness "
        "(standardized beta = 0.266, p < 0.001). By contrast, the broader Recovery "
        "Performance construct is weaker and more heterogeneous, and staffing/payroll "
        "intensity alone is not a robust positive correlate of it. Minority/language "
        "vulnerability is negatively associated with the performance construct, while "
        "socioeconomic and household composition/disability vulnerability are negatively "
        "associated with timeliness. These findings suggest that CDBG-DR administration "
        "depends not only on resource levels but on whether staffing is proportionate "
        "to program and disaster workload. Because the analysis pools jurisdictions "
        "across disaster cohorts and policy periods and relies on indirect staffing "
        "proxies, the results should be interpreted as exploratory pooled associations "
        "rather than causal effects."
    ),
    14: (
        "How is variation in state and local governmental capacity associated with "
        "disaster recovery program outcomes, and which institutional, socioeconomic, "
        "and programmatic conditions covary with recovery performance and timeliness?"
    ),
    15: (
        "How do secondary descriptive spatial summaries illustrate where workload "
        "manageability, recovery timeliness, and contextual vulnerability coincide "
        "across counties and state-administered recovery systems?"
    ),
    16: (
        "To address these questions, this study develops a cross-sectional Structural "
        "Equation Modeling (SEM) framework to conceptualize, measure, and compare "
        "state and local governmental capacity in administering CDBG-DR funds and "
        "implementing disaster recovery programs. Using administering-jurisdiction "
        "profiles constructed from quarterly QPR records and linked contextual "
        "covariates, the SEM is used as an exploratory latent-summary framework rather "
        "than as proof that every indicator is a fully validated reflective "
        "manifestation of a single underlying trait. This design does not estimate "
        "causal effects or formal moderation. Instead, state-versus-local status, "
        "jurisdiction size, and social vulnerability enter the SEM as observed "
        "covariates, while the spatial sections provide descriptive context for the "
        "core measurement and structural results."
    ),
    17: (
        "This approach advances the literature in three bounded ways. First, it "
        "operationalizes governmental capacity in CDBG-DR administration using both "
        "resource indicators and workload-manageability indicators rather than a single "
        "throughput ratio. Second, it compares a one-factor and a two-factor SEM so "
        "that the measurement tradeoffs are visible instead of being hidden behind a "
        "single preferred specification. Third, it treats the SEM as the article's core "
        "evidence and uses spatial summaries only as supplementary descriptive "
        "illustrations."
    ),
    23: (
        "In general, four broad approaches have been used to evaluate fund "
        "administration performance in disaster recovery. The first consists of "
        "descriptive and programmatic assessments. HUD and the Government "
        "Accountability Office, for example, document variation in fund utilization "
        "rates and administrative monitoring across grantees (U.S. Government "
        "Accountability Office, 2019; U.S. Department of Housing and Urban "
        "Development, 2026). While these studies reveal substantial variation in "
        "performance, they generally do not examine the institutional mechanisms that "
        "drive differences in disaster recovery outcomes. The second approach involves "
        "regression-based and panel-data analyses. Many studies in public "
        "administration and planning apply regression or longitudinal panel models to "
        "examine how staffing levels, administrative experience, governance "
        "structure, or disaster severity influence the pace of fund implementation, "
        "recovery outcomes, and broader disaster resilience (Cutter et al., 2016; "
        "Lee, 2019; Nor, 2025). These approaches provide useful empirical insight, "
        "but they typically treat financial indicators as directly observed outcomes "
        "and do not model governmental capacity as a latent construct reflected "
        "through multiple interrelated indicators."
    ),
    24: (
        "The third approach consists of case studies and qualitative evaluations. "
        "Qualitative research provides an in-depth understanding of administrative "
        "processes and identifies recurring implementation barriers, including "
        "procurement requirements, environmental review procedures, and coordination "
        "challenges across agencies (Jordan et al., 2014; Ospina et al., 2018; "
        "Rouhanizadeh et al., 2020). The fourth approach involves latent-variable "
        "methods, particularly Structural Equation Modeling (SEM). In the broader "
        "disaster research literature, SEM has been used to study multidimensional "
        "concepts such as preparedness, perceived risk, institutional resilience, and "
        "adaptive capacity (Geddam & Kiran, 2024; Gendeshmin et al., 2025). "
        "Applications of SEM to federal disaster-recovery finance remain limited, "
        "however, and few studies have used QPR data to construct latent measures of "
        "governmental capacity for CDBG-DR administration."
    ),
    27: (
        "Despite progress in understanding disaster recovery administration, several "
        "gaps remain. First, few studies develop multidimensional measures of "
        "governmental capacity that combine staffing, payroll, and workload-density "
        "indicators into one transparent framework. Second, cross-jurisdictional "
        "comparisons rarely distinguish resource intensity from workload "
        "manageability, even though those dimensions may behave differently. Third, "
        "the relationship between capacity and recovery outcomes is often examined "
        "without enough attention to maturity, right-censoring, and cross-cohort "
        "comparability. Fourth, spatial summaries are frequently presented without a "
        "clear statement of whether they are core evidence or descriptive context. "
        "These gaps motivate a more explicit and bounded SEM design."
    ),
    28: (
        "In response, this study develops a cross-sectional SEM of administering-"
        "jurisdiction capacity and recovery outcomes using linked QPR, QCEW, "
        "population, and SVI indicators. The article centers on the distinction "
        "between Administrative Resources and workload manageability and evaluates how "
        "those dimensions covary with Recovery Timeliness, while treating Recovery "
        "Performance as a broader and less settled secondary construct. The "
        "accompanying spatial summaries and gap screens are retained only as secondary "
        "descriptive context. Taken together, the paper offers a more transparent, "
        "policy-relevant, and explicitly exploratory framework for understanding "
        "variation in CDBG-DR administration across jurisdictions."
    ),
    31: (
        "The QPR dataset was obtained from the U.S. Department of Housing and Urban "
        "Development's DRGR reporting system and contains quarterly information on "
        "obligated, disbursed, and expended CDBG-DR funds for grants and reporting "
        "activities. The records span 18 disaster contexts from 2003 through 2023 and "
        "include thousands of organization labels that appear in grantee and activity "
        "fields (U.S. Department of Housing and Urban Development, 2026). Quarterly "
        "records were used to construct study-window outcome indicators for each "
        "administering jurisdiction, including the disbursed-to-obligated ratio, the "
        "expended-to-disbursed ratio, the share of obligated funds fully expended, the "
        "proportion of programs completed, the duration of grant completion, and the "
        "average duration of program completion. Throughout the manuscript, grantee "
        "refers to HUD's DRGR administrative label, activity responsible organization "
        "refers to the raw quarterly reporting name, and administering jurisdiction "
        "refers to the analytic SEM unit after those reporting entities are linked to "
        "a common state-agency or county-linked local-government profile. Grant "
        "denotes the disaster-specific CDBG-DR allocation, whereas activity denotes "
        "the underlying DRGR/QPR reporting unit nested within a grantee's disaster "
        "portfolio. The quarter-varying spending indicators therefore summarize the "
        "observed reporting history of each administering jurisdiction rather than a "
        "single quarter snapshot."
    ),
    32: (
        "Administrative resource data were drawn from the U.S. Bureau of Labor "
        "Statistics Quarterly Census of Employment and Wages (QCEW). Employment and "
        "payroll measures tied to NAICS 925110 (Administration of Housing Programs) "
        "are used as indirect proxies for disaster-recovery administrative resources "
        "(U.S. Bureau of Labor Statistics, 2026). These measures provide a "
        "standardized national frame, but they should not be interpreted as direct "
        "counts of staff assigned specifically to CDBG-DR. Disaster recovery "
        "administration may be distributed across agencies, temporary teams, contract "
        "support, and cross-departmental arrangements that are only partially captured "
        "by this industry code. Accordingly, employment and payroll are used as coarse "
        "administrative-resource indicators rather than exact staffing measures."
    ),
    33: (
        "Finally, contextual measures of social vulnerability were derived from the "
        "CDC/ATSDR Social Vulnerability Index (SVI) at the state and county levels. "
        "The analysis uses the four standard SVI themes - socioeconomic status, "
        "household composition and disability, minority status and language, and "
        "housing type and transportation - as broad contextual covariates rather than "
        "as precise longitudinal trend measures (CDC/ATSDR, 2026). Because CDC warns "
        "against literal cross-vintage comparison of SVI percentiles without careful "
        "harmonization, these variables are interpreted as standardized contextual "
        "descriptors in the final SEM-ready dataset rather than as year-by-year "
        "measures of change."
    ),
    36: (
        "Third, a new geographic identifier was constructed to determine the county and "
        "state associated with each Activity Responsible Org. The matching workflow "
        "followed a hierarchical approach: explicit county or parish names embedded in "
        "organization titles were matched first, and remaining records were linked "
        "through city-to-county lookup tables and manual review of unresolved cases. "
        "City-level entities were then aggregated to counties, while statewide "
        "departments, agencies, authorities, and similar organizations were aggregated "
        "to state-agency profiles. In practical terms, the geography workflow maps raw "
        "organization labels to one of three categories: state-agency profile, "
        "county-linked local-government profile, or unresolved case requiring manual "
        "judgment. This procedure provided a workable geography layer for descriptive "
        "mapping, but it should not be treated as a fully audited Census "
        "relationship-file reconstruction. Geography-dependent findings are therefore "
        "interpreted as informative spatial summaries rather than as exact "
        "administrative-boundary measures."
    ),
    37: (
        "Finally, a quality-control framework was implemented to evaluate the "
        "reliability of the inferred geographic information. Each record was categorized "
        "by match type (for example, explicit county match, city-to-county match, or "
        "unresolved) and flagged for manual review when automated inference was "
        "uncertain. Records with fewer than four quarterly observations were excluded "
        "from the cross-sectional SEM input because such short reporting histories do "
        "not provide enough information to summarize recovery performance or timeliness "
        "reliably. Appendix Table A4 reports the recoverable data flow used here: "
        "130,605 standardized QPR rows, 128,382 rows with valid quarter labels, 3,623 "
        "collapsed grantee-disaster-quarter records, and finally 573 administering-"
        "jurisdiction profiles - 30 state agencies and 543 local governments - for the "
        "complete-case SEM sample. If q indexes valid quarters for jurisdiction j, the "
        "analytical row stores study-window summaries Y_j = s({Y_jq}) rather than a "
        "single quarter value; if a indexes programs or activities nested within j, "
        "average program-duration measures summarize {D_ja}. The SEM therefore "
        "estimates between-jurisdiction differences in accumulated administrative "
        "profiles, not within-jurisdiction change over time. Appendix Table A8 reports "
        "the official 2023 Census crosswalk audit showing that all state-agency and "
        "county-linked SEM labels resolve to state or county GEOIDs."
    ),
    40: (
        "This study employs structural equation modeling (SEM) to develop latent "
        "constructs representing governmental capacity and disaster recovery outcomes "
        "and to evaluate the cross-sectional associations among them. The design does "
        "not estimate moderation effects. Instead, the state-versus-local "
        "indicator, jurisdiction size, and social vulnerability measures enter the SEM "
        "as observed covariates, and any state/local contrasts discussed later in the "
        "paper are treated as descriptive subset comparisons rather than as formal "
        "interaction tests. SEM remains useful in this setting as a structured way to "
        "summarize multiple indicators measured on different scales, but the pooled "
        "measurement model should be read as an exploratory approximation rather than as "
        "a demonstrated invariance result across state and local cases."
    ),
    42: (
        "Building on that framework, the analysis uses a cross-sectional dataset of "
        "573 administering-jurisdiction profiles, including 30 state agencies and 543 "
        "local jurisdictions. For state cases, the observation is the state agency that "
        "administers one or more CDBG-DR grants. For local cases, the observation is a "
        "county-linked local-government profile formed by aggregating city entities and "
        "activity organizations to a common administering jurisdiction. Each "
        "observation therefore summarizes an administering jurisdiction's resource and "
        "recovery profile across the study window rather than representing a single "
        "quarter or a single disaster episode. This aggregation permits a national "
        "SEM, but it is also a strong modeling assumption because it compresses "
        "variation across disaster cohorts, federal policy periods, staffing "
        "arrangements, and portfolio complexity into one cross-sectional profile. "
        "Accordingly, county is used in the paper as the geographic container for "
        "local profiles, state agency is used for state-level profiles, and grantee is "
        "reserved for the original HUD label rather than as a synonym for the SEM's "
        "analytical unit."
    ),
    44: (
        "In this study, governmental capacity is conceptualized as the institutional "
        "ability of public agencies to administer CDBG-DR funds under differing levels "
        "of workload and contextual constraint. The two-factor specification separates "
        "Administrative Resources from workload manageability, which is labeled "
        "Administrative Burden Capacity in the SEM output. Administrative Resources are "
        "represented by staffing and payroll intensity proxies derived from QCEW-based "
        "employment and payroll measures, while workload manageability is represented "
        "by staff-scaled workload indicators that relate the number of disaster "
        "programs and disaster events to available staffing. This two-part "
        "conceptualization is used as an exploratory latent summary of administrative "
        "infrastructure and workload pressure rather than as definitive proof that the "
        "indicators are interchangeable reflective manifestations of a single capacity "
        "trait."
    ),
    45: (
        "Two additional indicators reflect administrative workload intensity relative "
        "to staffing levels: the number of disaster recovery programs per staff "
        "(programs_per_staff) and the number of disaster events per staff "
        "(disasters_per_staff). Higher workload per staff implies greater "
        "administrative strain, not greater capacity. To keep coefficient signs "
        "interpretable, both workload measures are reverse-coded so that higher values "
        "on the SEM input correspond to lower relative burden and, therefore, more "
        "favorable workload conditions. Operationally, "
        "programs_per_staff_j = NumProgram_j / (avg_employment_j x E_TOTPOP_j + e) and "
        "disasters_per_staff_j = NumDisaster_j / (avg_employment_j x E_TOTPOP_j + e), "
        "where avg_employment is a population-normalized staffing proxy, E_TOTPOP is "
        "jurisdiction population, and e is a small constant used only to avoid "
        "division-by-zero. Substantively, these indicators can be read as more staff "
        "capacity per program and per disaster."
    ),
    53: (
        "Disaster recovery outcomes are modeled as two latent constructs: Recovery "
        "Performance and Recovery Timeliness. Recovery Performance captures the "
        "effectiveness of financial management and program implementation and is "
        "measured using four indicators: the disbursed-to-obligated ratio, the "
        "expended-to-disbursed ratio, the share of obligated funds fully expended, and "
        "the proportion of programs completed. Together, these measures summarize how "
        "effectively jurisdictions move funds through the recovery pipeline and convert "
        "those funds into completed programs. These are cumulative administrative "
        "summary measures, not event-history outcomes. Because DRGR portfolios are "
        "updated through corrections, reimbursements, and accounting adjustments, some "
        "observed cumulative ratios can exceed 1.0 in the raw administrative file, "
        "especially for expended-to-disbursed and fully-expended measures. The main "
        "SEM retains those observed values rather than mechanically truncating them, "
        "and interprets them as imperfect but informative administrative indicators. "
        "Appendix Table A7 reports how often ratios exceed 1.0 and shows a capped-ratio "
        "sensitivity check so that this accounting feature is explicit rather than "
        "hidden."
    ),
    54: (
        "Recovery Timeliness represents the speed and consistency of program and grant "
        "implementation. It is measured using the duration of grant completion and the "
        "average duration of program completion, both reverse-coded so that higher "
        "values indicate faster implementation. These are useful summary indicators, "
        "but they are also sensitive to maturity and right-censoring. Newer grants and "
        "ongoing portfolios have had less time to close out, so timeliness is treated "
        "as a cross-sectional summary of observed completion pace rather than as a "
        "fully uncensored event-history outcome. Duration of completion records months "
        "to 95% expenditure for the jurisdictional portfolio, while average program "
        "duration summarizes completion time across completed program or activity "
        "components within that portfolio."
    ),
    59: (
        "To account for contextual differences across jurisdictions, the SEM includes "
        "governmental level, jurisdiction size, and social vulnerability covariates "
        "(Peacock et al., 2014; CDC/ATSDR, 2026). Governmental level is represented by "
        "a binary indicator coded 1 for state agencies and 0 for local governments. "
        "This variable is modeled as a direct predictor and later used to organize "
        "descriptive subset comparisons; it is not a moderation term. Jurisdiction size "
        "is captured through total population, and the four SVI themes are included as "
        "broad contextual controls. Because the final SEM-ready dataset stores these "
        "covariates in standardized form, coefficients reflect cross-sectional "
        "associations within the aggregated administering-jurisdiction sample rather "
        "than temporal effects of changing population or SVI conditions. State and "
        "local cases are therefore adjusted within a pooled model, but they are not "
        "treated as institutionally interchangeable units for strong comparative "
        "inference."
    ),
    61: (
        "The primary SEM specification decomposes governmental capacity into two "
        "related latent dimensions - Administrative Resources and Administrative Burden "
        "Capacity - while representing disaster recovery outcomes as two related latent "
        "constructs - Recovery Performance and Recovery Timeliness. This specification "
        "is used as the main interpretive model because it better separates resource "
        "levels from workload pressure and produces stronger conventional fit indices "
        "than the one-factor alternative, even though the alternative retains lower "
        "AIC and BIC. The two-factor SEM is therefore retained as the article's main "
        "summary model, while the underlying measurement logic remains explicitly "
        "exploratory."
    ),
    63: (
        "The SEM framework enables the analysis to distinguish between the effects of "
        "greater administrative resources and those associated with lower relative "
        "workload burdens. Both latent capacity dimensions are specified as predictors "
        "of Recovery Performance and Recovery Timeliness while controlling for "
        "governmental level, jurisdiction size, and the four SVI dimensions. The two "
        "capacity dimensions are allowed to covary, as are the two outcome constructs. "
        "These are modeled as direct relationships rather than as interaction or "
        "moderation effects."
    ),
    74: (
        "The working expectation is that jurisdictions with greater administrative "
        "resources and lower relative workload pressure will exhibit stronger recovery "
        "outcomes. This expectation is treated as a cross-sectional associational "
        "hypothesis rather than as a claim of causal identification."
    ),
    76: (
        "Two model-comparison and sensitivity strategies were implemented. First, the "
        "analysis compares a two-factor SEM, in which Administrative Resources and "
        "Administrative Burden Capacity are estimated separately, with a one-factor SEM "
        "that combines all four indicators into a single Governmental Capacity factor. "
        "Second, archived references to a smaller cleaned sample (N = 169) are retained "
        "only for provenance documentation. Because the exact filtering rule for that "
        "subset cannot be reconstructed, it is no longer used as a substantive "
        "robustness test for the paper's claims. Appendix Table A10 shows that the "
        "smaller sample is a strict subset of the 573-jurisdiction file but does not "
        "appear to be recoverable from any simple one-threshold outlier screen."
    ),
    80: (
        "Prior to SEM estimation, all continuous indicators were standardized as "
        "z-scores to improve comparability across variables and to support model "
        "convergence, while the state_level indicator remained coded 0/1. The "
        "complete-case SEM input contains 573 jurisdictions and 17 variables with no "
        "remaining missing values. Workload measures use the staff-scaled "
        "denominators described above, and a small constant is used during ratio "
        "construction to avoid division-by-zero problems in jurisdictions with very "
        "small staffing counts. The SEMs were estimated with conventional maximum "
        "likelihood estimation (Lee & Zhu, 2002), so the fit statistics, standard "
        "errors, confidence intervals, and p-values reported here are conventional ML "
        "quantities. Model "
        "evaluation relies on a combination of fit statistics - chi-square, CFI, "
        "TLI, RMSEA, AIC, and BIC - rather than on a single threshold. Model "
        "comparison is therefore treated as a tradeoff exercise: the two-factor model "
        "is substantively more informative and fits better on the main incremental "
        "and absolute fit measures, while the one-factor model remains more "
        "parsimonious by information criteria. Appendix Table A3 verifies the archived "
        "fit statistics against reruns on the recovered 573-jurisdiction SEM-ready "
        "dataset, including the reported AIC and BIC values."
    ),
    84: (
        "Both SEM specifications were estimated to examine the relationship between "
        "governmental capacity and disaster recovery outcomes: (1) a one-factor model "
        "in which all four capacity indicators load on a single Governmental Capacity "
        "factor and (2) a two-factor model that separates Administrative Resources from "
        "Administrative Burden Capacity. Table 1 reports the comparison for the "
        "573-jurisdiction complete-case SEM sample."
    ),
    86: (
        "The model comparison reveals a genuine tradeoff rather than an unambiguous "
        "winner. The one-factor model has lower information criteria (AIC = 67.41; "
        "BIC = 219.69), but the two-factor model fits the covariance structure more "
        "convincingly on the conventional SEM diagnostics (chi-square(98) = 469.82 "
        "versus chi-square(101) = 741.72; both chi-square p-values < 0.001; CFI = "
        "0.915 versus 0.854; TLI = 0.891 versus 0.818; RMSEA = 0.081 versus 0.105). "
        "Because the two-factor model aligns more closely with the theory that "
        "resources and workload are distinct dimensions of administrative capacity, "
        "and because it improves the main fit indices materially, it is used as the "
        "primary substantive model in the remainder of the paper. At the same time, "
        "the one-factor alternative is more parsimonious by AIC/BIC, so the two-factor "
        "model should be interpreted as preferred conditionally rather than as "
        "definitively superior on every criterion. Appendix Table A3 reports the "
        "imported-versus-rerun fit comparison and confirms that the unusual AIC/BIC "
        "scale reflects the archived output convention rather than a transcription "
        "error."
    ),
    87: (
        "The measurement parameters of the two-factor model also motivate a cautious "
        "but interpretable reading. Within Administrative Resources, standardized "
        "loadings are strong for employment (0.953) and payroll (0.992). Within "
        "Administrative Burden Capacity, the reverse-coded programs-per-staff "
        "indicator loads at 0.577 and disasters-per-staff loads at approximately "
        "1.000, indicating that the second factor is driven more heavily by the "
        "disaster-workload component. Recovery Performance is less coherent: "
        "disbursed-to-obligated loads at 0.785, obligated funds fully expended at "
        "0.988, programs completed at 0.477, and expended-to-disbursed at -0.131. "
        "Recovery Timeliness is identified by reverse-coded duration of completion "
        "(0.545) and reverse-coded average program duration (1.000). The two latent "
        "capacity dimensions are positively correlated (standardized covariance = "
        "0.402), and the structural part of the model explains more variance in "
        "Recovery Timeliness than in Recovery Performance (approximate latent R^2 = "
        "0.20 versus 0.09, based on standardized residual variances). These features do "
        "not invalidate the two-factor model, but they do argue for interpreting it as "
        "an exploratory cross-sectional measurement structure rather than as a final, "
        "fully settled confirmatory SEM. Recovery Timeliness is the more coherent and "
        "policy-relevant outcome in this specification, whereas Recovery Performance is "
        "better treated as a broader secondary summary of financial and programmatic "
        "execution. Appendix Table A1 reports the full recovered measurement summary, "
        "Appendix Table A2 reports the complete structural and covariance output, "
        "Appendix Table A6 shows the main sensitivity checks, and Appendix Table A9 "
        "summarizes the proxy-validation diagnostics."
    ),
    89: (
        "The structural results indicate that Administrative Resources and "
        "Administrative Burden Capacity relate to different recovery outcomes in "
        "different ways. Table 2 reports a highlighted subset of the standardized "
        "structural path coefficients from the two-factor SEM together with "
        "unstandardized 95% confidence intervals. Appendix Table A2 reports the full "
        "modeled path set, factor covariances, and residual terms. Supplementary "
        "SEM-implied relationship plots and a limited set of descriptive county-level "
        "maps are retained only as secondary illustrations; the core evidentiary weight "
        "of the paper rests on the SEM and its measurement and structural parameters."
    ),
    91: (
        "Note: Table 2 reports a highlighted subset of standardized coefficients from "
        "the two-factor SEM. The final column reports 95% confidence intervals for the "
        "corresponding unstandardized path coefficient together with exact p-values. "
        "Appendix Table A2 reports the complete modeled coefficient set."
    ),
    93: (
        "Administrative Resources exhibit a negative association with Recovery "
        "Performance (beta = -0.170; unstandardized b = -0.139, 95% CI [-0.212, "
        "-0.066]; p < 0.001), indicating that jurisdictions with higher staffing and "
        "payroll intensity do not necessarily display stronger financial execution or "
        "program completion once workload pressure and contextual conditions are taken "
        "into account. This coefficient should not be read as evidence that more "
        "resources worsen recovery; at most it suggests that resource intensity may "
        "proxy for more complex portfolios or administrative settings that the pooled "
        "cross-sectional model does not fully observe. By contrast, Administrative "
        "Burden Capacity is positive but only marginally associated with Recovery "
        "Performance (beta = 0.077; b = 0.104, 95% CI [-0.015, 0.222]; p = 0.086). "
        "Manageable workload appears helpful for performance, but the evidence is much "
        "stronger for timeliness than for the broader Recovery Performance construct."
    ),
    94: (
        "Among the covariates, SVI Theme 3 (minority status and language vulnerability) "
        "shows the clearest association with Recovery Performance (beta = -0.226, "
        "p < 0.001). The predicted relationship plot in Figure 2(a) indicates that "
        "Recovery Performance tends to decline as minority and language vulnerability "
        "increase. This pattern should be read as contextual rather than causal: "
        "jurisdictions serving more linguistically diverse or historically marginalized "
        "communities may face greater outreach, translation, verification, and service-"
        "delivery demands, not because those communities cause weaker recovery, but "
        "because administrative barriers and structural inequities can be harder to "
        "overcome in those settings."
    ),
    96: (
        "Figure 2. SEM-implied predicted relationship: (a) SVI Theme 3: minority "
        "status and language → Recovery Performance, (b) SVI Theme 1: socioeconomic "
        "→ Recovery Timeliness, and (c) SVI Theme 2: household composition/disability "
        "→ Recovery Timeliness."
    ),
    98: (
        "For Recovery Timeliness, Administrative Burden Capacity plays the clearest "
        "role. Lower workload pressure relative to available staffing is strongly "
        "associated with faster and more consistent recovery implementation "
        "(beta = 0.266; b = 0.251, 95% CI [0.151, 0.352]; p < 0.001), indicating that "
        "administrative strain is an important constraint on implementation speed. "
        "Administrative Resources are positively associated with Recovery Timeliness "
        "as well, but only marginally (beta = 0.080; b = 0.046, 95% CI [-0.004, "
        "0.095]; p = 0.072). This contrast reinforces the main substantive message of "
        "the paper: resource levels matter, but workload manageability appears to be "
        "the more consistent cross-sectional correlate of timely recovery "
        "administration."
    ),
    99: (
        "Two dimensions of social vulnerability are also significantly associated with "
        "Recovery Timeliness. Higher values of SVI Theme 1 (socioeconomic "
        "vulnerability) and SVI Theme 2 (household composition and disability) are "
        "associated with slower implementation (beta = -0.194, p = 0.004; beta = "
        "-0.159, p = 0.004). Figures 2(b) and 2(c) show these downward patterns. The "
        "results suggest that jurisdictions serving more socioeconomically vulnerable "
        "and demographically constrained populations may face additional administrative "
        "frictions that slow recovery even after capacity proxies are taken into "
        "account."
    ),
    100: (
        "The state-versus-local indicator is not statistically meaningful for Recovery "
        "Performance (beta = -0.038; b = -0.133, 95% CI [-0.526, 0.259]; p = 0.505) "
        "and is only borderline negative for Recovery Timeliness (beta = -0.106; b = "
        "-0.260, 95% CI [-0.520, 0.000]; p = 0.050). Governmental level is therefore "
        "treated as a control variable and a descriptive comparison device, not as a "
        "moderator of capacity effects. The negative sign for timeliness suggests that "
        "local governments may, on average, move somewhat faster than state agencies "
        "once other covariates are controlled, but the evidence is too modest and the "
        "institutional tasks too heterogeneous to support strong claims about "
        "institutional superiority."
    ),
    101: "5.2.3. Supplementary SEM-implied relationship plots",
    102: (
        "The SEM-implied predicted relationship plots further indicate that "
        "Administrative Burden Capacity explains more of the visible variation in "
        "recovery outcomes than Administrative Resources alone. Figures 3(a) and 3(b) "
        "show upward associations between Administrative Burden Capacity and both "
        "Recovery Performance and Recovery Timeliness, with the relationship appearing "
        "substantially stronger for timeliness. These plots are useful visualization "
        "tools, but they should not be interpreted as evidence of nonlinear causal "
        "effects. They are derived from the fitted SEM and are therefore best read as "
        "model-implied association curves."
    ),
    105: (
        "Figure 3. SEM-implied predicted relationship: (a) Administrative Burden "
        "Capacity -> Recovery Performance, (b) Administrative Burden Capacity -> "
        "Recovery Timeliness, (c) Administrative Resources -> Recovery Performance, and "
        "(d) Administrative Resources -> Recovery Timeliness."
    ),
    106: (
        "Administrative Resources exhibit a more complex relationship with recovery "
        "outcomes. Figure 3(c) mirrors the negative coefficient reported in Table 2 "
        "for Recovery Performance, whereas Figure 3(d) reflects the weaker positive "
        "association with Recovery Timeliness. One plausible interpretation is that "
        "larger staffing and payroll footprints are characteristic of more complex or "
        "higher-demand recovery portfolios. Once workload is modeled explicitly, the "
        "resource indicators no longer behave like simple monotonic measures of "
        "administrative effectiveness."
    ),
    108: (
        "Supplementary indicator-level predicted relationship plots were also "
        "generated using the observed variables that underlie the latent capacity "
        "constructs. These figures are presented in Appendix Figures A1-A5. They are "
        "retained as descriptive visual checks rather than as separate evidentiary "
        "tests, and they broadly support the same interpretation: staff-scaled "
        "workload measures align more consistently with Recovery Timeliness than raw "
        "resource intensity alone."
    ),
    110: (
        "To complement the SEM results, the manuscript retains only a small "
        "appendix-style descriptive spatial section showing how selected county-level "
        "capacity and recovery scores are distributed geographically. These figures "
        "are not part of the SEM estimation, and they are not used to establish "
        "additional statistical claims. Instead, they provide visual context for "
        "whether the burden-capacity pattern identified in the SEM appears "
        "concentrated in particular regions."
    ),
    111: (
        "Appendix Table A8 summarizes the official 2023 Census crosswalk audit used "
        "for those figures by distinguishing state-agency profiles, county-linked "
        "local profiles, and any unresolved records. In the recovered 573-jurisdiction "
        "sample used here, the archived labels resolve cleanly to official state or "
        "county GEOIDs, but that result should still be interpreted as an audit of the "
        "recovered sample labels rather than a full reconstruction of every earlier "
        "hierarchical matching decision."
    ),
    112: (
        "The descriptive maps are kept for orientation rather than inference. They "
        "summarize the geographic distribution of latent scores and related overlays, "
        "but they are not estimated as part of the SEM and should not be read as "
        "independent tests of the article's theoretical claims."
    ),
    115: (
        "Across the retained maps, the main descriptive pattern is straightforward: "
        "workload manageability and Recovery Timeliness appear more closely aligned "
        "than raw Administrative Resources and Recovery Performance. Because these "
        "maps inherit values from county-linked and state-agency profiles, they are "
        "best interpreted as secondary descriptive context."
    ),
    116: (
        "For that reason, the map discussion is intentionally concise. The maps are "
        "used only to show whether the burden-capacity pattern identified in the SEM "
        "has a visible geographic footprint, not to establish an additional spatial "
        "model."
    ),
    117: "5.3.2. Supplementary appendix-style descriptive overlays",
    118: (
        "The quadrant overlays and related county-level summaries are retained only as "
        "supplementary illustrations. Their main value is to show that the stronger "
        "burden-timeliness relationship does not appear purely abstract when mapped "
        "descriptively."
    ),
    119: (
        "Appendix Figure A6. Descriptive quadrant overlays of standardized "
        "Administrative Burden Capacity with (a) Recovery Timeliness and (b) Recovery "
        "Performance at the county level."
    ),
    120: (
        "Appendix Figure A6(a) is most useful as a visual cross-check on the SEM's main "
        "substantive result: counties with weaker burden-capacity scores are more "
        "often found in the same descriptive space as slower Recovery Timeliness. The "
        "map does not establish a spatial effect, but it helps show that the burden-"
        "timeliness relationship is not purely abstract."
    ),
    121: (
        "Appendix Figure A6(b) shows a noticeably weaker descriptive relationship for Recovery "
        "Performance, which is consistent with the SEM result that the performance "
        "construct is more heterogeneous and less tightly linked to burden capacity "
        "than Recovery Timeliness is."
    ),
    122: (
        "Overall, the limited spatial summaries are consistent with the paper's "
        "central descriptive conclusion that workload manageability appears more "
        "closely aligned with Recovery Timeliness than raw Administrative Resources "
        "alone. Because the geography layer relies on hierarchical name matching and "
        "county aggregation, the maps should be interpreted as descriptive context "
        "rather than as precise jurisdictional rankings."
    ),
    124: (
        "The article also retains exploratory gap diagnostics and a governance-risk "
        "screen in the appendix. These are heuristic descriptive overlays, not latent "
        "variables or causal estimands, and they are not used as principal findings."
    ),
    126: (
        "Among those appendix-only overlays, the administrative burden-capacity gap is "
        "the most interpretable because it asks whether observed exposure is high "
        "relative to workload manageability. Even there, the measure is best treated as "
        "a preliminary screen for potential administrative strain rather than as a "
        "standalone index of governance quality."
    ),
    128: (
        "The exposure-capacity typology therefore serves only as a descriptive screen "
        "for jurisdictions where exposure appears high relative to workload "
        "manageability. It should not be read as a validated upstream measure of "
        "governance quality."
    ),
    130: (
        "The implementation gap and performance gap are retained only for descriptive "
        "completeness and are not needed to understand the paper's main SEM result."
    ),
    132: (
        "The exposure-timeliness typology is likewise retained only as a descriptive "
        "overlay of observed strain. It may help practitioners identify places where "
        "implementation appears slower than expected given exposure, but it is not a "
        "separate causal or spatial model."
    ),
    134: (
        "Accordingly, the gap maps are interpreted only as descriptive overlays that "
        "suggest where high exposure, weaker burden capacity, and weaker outcomes may "
        "coincide. They are not used to adjudicate the main theoretical argument of "
        "the paper."
    ),
    136: (
        "The exposure-performance typology adds the least independent information of "
        "the three gap screens because it repackages the same performance construct "
        "already modeled in the SEM. It is therefore retained only for descriptive "
        "completeness in the appendix."
    ),
    137: (
        "Taken together, these exploratory difference scores point to a similar broad "
        "geographic pattern as the SEM and the simpler spatial overlays, but they do "
        "not materially strengthen the core argument. Their value is screening and "
        "illustration, not independent confirmation."
    ),
    139: (
        "The Recovery Governance Risk Index is likewise presented only as a heuristic "
        "descriptive composite. It increases with disaster exposure and social "
        "vulnerability and decreases with burden capacity, but it is not estimated as "
        "part of the SEM and it should not be treated as a validated ranking model."
    ),
    140: (
        "Appendix Figure A7. Supplementary descriptive screens of (a) county-level "
        "disaster recovery governance risk and (b) high-governance-risk counties with "
        "elevated descriptive priority."
    ),
    141: (
        "Like the gap diagnostics, the risk map points descriptively toward parts of "
        "the Gulf Coast and South, but it is retained as an exploratory overlay rather "
        "than as a principal empirical result."
    ),
    142: (
        "The thresholded priority map is therefore best read as an informal triage aid "
        "rather than as a publication-grade league table."
    ),
    143: (
        "Its value, if used at all, lies in pointing practitioners toward places where "
        "more careful institutional diagnostics may be warranted. It is not a central "
        "contribution of the article."
    ),
    144: (
        "Appendix Tables A1-A10 report the supporting measurement, full structural, "
        "fit-verification, data-flow, maturity, sensitivity, ratio-artifact, "
        "geography, proxy-validation, and subset-forensics summaries that accompany "
        "the core SEM."
    ),
    147: (
        "This study advances a more differentiated understanding of governmental "
        "capacity in disaster recovery by showing that capacity is not well captured by "
        "a single resource proxy. Across the SEM and the limited descriptive spatial "
        "overlays retained here, the most consistent pattern is that recovery outcomes are "
        "more closely aligned with workload manageability than with staffing and "
        "payroll intensity alone. That does not mean resources are irrelevant; rather, "
        "it suggests that resources must be interpreted relative to the disaster and "
        "program burdens they are expected to absorb."
    ),
    149: (
        "A central contribution of the paper is the distinction between "
        "Administrative Resources and Administrative Burden Capacity. Staffing and "
        "payroll reflect the scale of the administrative system, whereas burden "
        "capacity reflects how stretched that system is across disasters and programs. "
        "The empirical evidence is strongest for Recovery Timeliness, where lower "
        "workload pressure is associated with faster implementation. For Recovery "
        "Performance, the burden-capacity path is positive but only marginal, which "
        "suggests that the performance construct also reflects other dimensions of "
        "organizational context, program design, and reporting complexity. Burden "
        "capacity appears especially important for timing, and plausibly relevant for "
        "performance, but not sufficient by itself to explain all observed differences "
        "in recovery execution."
    ),
    150: (
        "The negative association between Administrative Resources and Recovery "
        "Performance further supports this interpretation. Rather than implying that "
        "additional staff or payroll harms recovery, the coefficient likely reflects "
        "portfolio complexity, organizational layering, or the fact that more "
        "administratively intensive jurisdictions often manage more demanding recovery "
        "contexts. Once workload is modeled explicitly, pure resource intensity ceases "
        "to function as a simple proxy for effectiveness. This result is therefore best "
        "treated as a potentially confounded association that cautions against equating "
        "larger bureaucratic scale with stronger recovery administration."
    ),
    152: (
        "The results also show that different dimensions of social vulnerability "
        "relate to recovery outcomes through distinct channels. Minority "
        "status and language vulnerability are most clearly associated with weaker "
        "Recovery Performance, whereas socioeconomic vulnerability and household "
        "composition/disability are more strongly associated with slower Recovery "
        "Timeliness. One plausible interpretation is that vulnerability marks contexts "
        "in which outreach, verification, coordination, and service delivery barriers "
        "are harder to overcome. The SEM does not identify those mechanisms directly, "
        "but the differentiated theme coefficients are consistent with the idea that "
        "recovery governance becomes more difficult when administrative burden and "
        "structural vulnerability overlap."
    ),
    153: (
        "The limited spatial overlays are retained only as supplementary screening "
        "tools. They may help show where workload pressure, exposure, and vulnerability "
        "coincide, but they do not establish why those patterns arise."
    ),
    154: (
        "For that reason, the maps are used descriptively rather than as a separate "
        "source of causal or confirmatory evidence."
    ),
    155: (
        "The article's main contribution therefore remains the burden-sensitive SEM "
        "framework, not the broader atlas of descriptive spatial screens."
    ),
    157: (
        "The findings provide decision-makers and practitioners with a more nuanced "
        "understanding of governmental capacity by distinguishing resource levels from "
        "workload strain. For practice, the key implication is diagnostic: when "
        "recovery performance is weak or implementation is slow, the problem may not be "
        "simple resource scarcity. It may instead reflect the way staffing, portfolio "
        "complexity, disaster exposure, and community vulnerability interact. This "
        "perspective can help practitioners diagnose whether recovery challenges stem "
        "from insufficient resources, excessive administrative burden, or a broader "
        "misalignment between institutional capacity and local recovery needs."
    ),
    158: (
        "Any use of the Recovery Governance Risk Index or the related gap diagnostics "
        "should therefore remain cautious and secondary. At most, these descriptive "
        "screens can help identify jurisdictions that may merit closer institutional "
        "review or technical assistance; they are not the article's principal empirical "
        "contribution and are best read in the appendix or supplement."
    ),
    159: (
        "From a policy perspective, the results suggest that improving disaster "
        "recovery outcomes requires moving beyond resource-centric approaches "
        "toward workload-sensitive and context-aware interventions. Simply increasing "
        "funding or staffing may have limited effect if administrative systems remain "
        "overextended or if program complexity continues to impose substantial "
        "procedural burden. More promising strategies may include simplifying program "
        "rules, standardizing reporting and eligibility workflows, deploying temporary "
        "surge-capacity teams during high-demand periods, improving intergovernmental "
        "coordination, and prioritizing high-risk jurisdictions for language access, "
        "community outreach, and implementation support. These recommendations remain "
        "associationally motivated rather than experimentally verified, but they follow "
        "directly from the burden-sensitive interpretation of the results."
    ),
    160: (
        "Although the real-world effectiveness of these strategies requires additional "
        "evaluation, the findings suggest that successful disaster recovery "
        "depends not only on the amount of resources available but also on whether "
        "those resources are adequate for the complexity, scale, and vulnerability of "
        "the recovery environment in which they are deployed."
    ),
    162: (
        "Several limitations should be acknowledged. First, the SEM relies on a "
        "cross-sectional administering-jurisdiction dataset, so the results are "
        "associational rather than causal. The aggregation from quarterly QPR activity "
        "data to one profile per administering jurisdiction discards "
        "within-jurisdiction variation across disasters, policy regimes, staffing "
        "arrangements, and maturation windows. Appendix Table A5 makes that maturity "
        "problem visible by showing that the sample spans markedly different portfolio "
        "age bands. Second, the timeliness construct "
        "remains duration-based and therefore sensitive to right-censoring and "
        "disaster maturity. Newer grants and jurisdictions with ongoing portfolios "
        "have had less time to complete recovery activities, which limits how strongly "
        "the timeliness factor can be interpreted."
    ),
    163: (
        "Third, the measurement of capacity remains indirect. QCEW NAICS 925110 "
        "employment and payroll are useful national proxies but do not directly measure "
        "CDBG-DR administrative teams, contractors, or interagency support structures. "
        "The workload indicators are likewise simplified workload-density proxies "
        "rather than complete measures of administrative strain. Appendix Table A6 "
        "shows that the main burden-timeliness result survives several conservative "
        "re-specifications, Appendix Table A7 reports the frequency of ratio values "
        "above 1.0 and a capped-ratio sensitivity, and Appendix Table A9 adds proxy-"
        "correlation and coarse external cross-check diagnostics, but these remain "
        "partial validation steps "
        "rather than definitive proof that the archived QCEW measures isolate true "
        "CDBG-DR staffing capacity. "
        "Similarly, the SVI covariates are used as cross-sectional contextual controls "
        "rather than as audited longitudinal trajectories across changing vintages. "
        "Fourth, the geography layer is derived from hierarchical name matching and "
        "county aggregation rather than from official Census relationship files alone, "
        "even though the recovered SEM sample labels now audit cleanly against "
        "official 2023 Census state and county files. Fifth, the smaller cleaned "
        "sensitivity sample remains only partially reconstructable; Appendix Table A10 "
        "shows that it behaves like a compound or curated filter rather than a simple "
        "outlier trim. Finally, the gap diagnostics and governance risk index are "
        "descriptive composites rather than validated latent constructs. For these "
        "reasons, the spatial and gap-based sections should be read as informative "
        "policy screens rather than as final measurement products."
    ),
    164: (
        "These limitations point to a clear future research agenda. Longitudinal or "
        "event-history designs would address maturity and censoring more directly. "
        "More detailed administrative-process data - such as staffing assignments, "
        "contractor use, application processing times, and intergovernmental "
        "coordination mechanisms - would make it possible to move from proxy capacity "
        "measures toward more direct operational indicators of recovery governance. "
        "Future work should also test narrower cohorts or panel designs so that age of "
        "portfolio and disaster vintage are separated more clearly from administrative "
        "capacity."
    ),
    165: (
        "Despite these limitations, the study still provides a useful empirical and "
        "conceptual contribution. Across the SEM and the descriptive spatial "
        "screens, the most consistent message is that disaster recovery governance is "
        "shaped not simply by how many resources a jurisdiction has, but by whether "
        "those resources are proportionate to disaster workload and to the social "
        "conditions of implementation. That is a meaningful insight even when treated "
        "explicitly as exploratory rather than causal."
    ),
    167: (
        "This study develops an explicitly exploratory framework for examining "
        "governmental capacity and disaster recovery outcomes across the United States. "
        "Using pooled cross-sectional administering-jurisdiction profiles, the article "
        "moves beyond purely resource-based indicators and evaluates how staffing "
        "proxies, workload pressure, and contextual vulnerability co-occur with "
        "Recovery Timeliness and, more tentatively, Recovery Performance. The SEM is "
        "the paper's core evidence, while the gap diagnostics and governance risk index "
        "are treated as supplementary policy screens."
    ),
    168: (
        "Several key insights emerge. First, governmental capacity is not well captured "
        "as a unidimensional construct. The two-factor model shows that Administrative "
        "Resources and Administrative Burden Capacity behave differently and should not "
        "be collapsed without checking the measurement tradeoff. Second, the strongest "
        "capacity-outcome association in the SEM is between lower workload "
        "pressure and faster Recovery Timeliness, whereas the relationship with the "
        "broader Recovery Performance construct is weaker and more conditional. Third, "
        "social vulnerability matters, but different SVI themes relate to different "
        "parts of the recovery process, reinforcing the view that capacity is embedded "
        "in broader institutional and community contexts."
    ),
    169: (
        "The limited descriptive spatial overlays add a geographic dimension to those "
        "findings, but only in a supplementary sense. They help orient the reader to "
        "where exposure, workload pressure, vulnerability, and weaker outcomes may "
        "overlap, while remaining clearly secondary to the SEM."
    ),
    170: (
        "Taken together, the study contributes conceptually by advancing a "
        "workload-sensitive view of governmental capacity, methodologically by making "
        "the one-factor versus two-factor SEM tradeoff explicit, and practically by "
        "showing that workload-sensitive diagnostic screening may be more useful than "
        "simple staffing counts alone. The paper should therefore be read as an "
        "exploratory pooled cross-sectional SEM study, centered on workload "
        "manageability and Recovery Timeliness, rather than as a definitive causal "
        "account of CDBG-DR recovery governance."
    ),
    171: (
        "Improving disaster recovery governance will require more than funding alone. "
        "It will require institutions whose staffing, workload, and social context are "
        "better aligned with the scale and complexity of recovery needs. The results "
        "provide an empirical starting point for that discussion."
    ),
    175: (
        "The authors declared no potential conflicts of interest with respect to the "
        "research, authorship, and/or publication of this article. Generative AI tools "
        "were used only for limited language editing. All analytic decisions, "
        "interpretation, and manuscript content were reviewed and approved by the "
        "authors."
    ),
    177: (
        "The public source datasets used in this study are HUD CDBG-DR DRGR/QPR "
        "reports, U.S. Bureau of Labor Statistics QCEW files, and CDC/ATSDR SVI data. "
        "The replication package for the published manuscript will include the code, "
        "SEM-ready analytic data, derived geography linkage files, and nonrestricted "
        "outputs needed to reproduce the reported models, tables, and maps. Public "
        "source data will be cited directly, and any remaining manually audited linkage "
        "notes will be provided as supplemental derived documentation."
    ),
    186: "",
    184: (
        "Costa, R., Mann, B., Sobhani, A., Hamideh, S., Nejat, A., & Ross, A. (2026). "
        "From Allocation to Action: A Comparative Analysis of CDBG-DR Funding "
        "Expenditures for 2017 Disasters in California, Florida, and Texas [working "
        "paper / SSRN preprint]. http://dx.doi.org/10.2139/ssrn.6153946"
    ),
    189: "",
    190: "",
    191: "",
    192: (
        "Gigasheet. (2026). List of All U.S. Counties and Cities Spreadsheet "
        "(public lookup table used for preliminary city-to-county matching). "
        "https://www.gigasheet.com/sample-data/list-of-all-us-counties-and-cities-spreadsheet"
    ),
    193: "",
    226: "",
    227: "",
    228: "",
    233: "",
    235: "",
    236: "",
    238: "",
    241: (
        "Figure A1. Indicator-level predicted relationship plots between programs per "
        "staff and recovery timeliness."
    ),
    242: (
        "Figure A2. Indicator-level predicted relationship plots between programs per "
        "staff and recovery performance."
    ),
    243: (
        "Figure A3. Indicator-level predicted relationship plots between disasters per "
        "staff and recovery performance."
    ),
    244: (
        "Figure A4. Indicator-level predicted relationship plots between average "
        "payroll and recovery performance."
    ),
    246: (
        "Figure A5. Indicator-level predicted relationship plots between average "
        "employment and recovery performance."
    ),
}

FORBIDDEN_STANDALONE_PHRASES = (
    "revised manuscript",
    "revised paper",
    "revised study",
    "revised analysis",
    "revised design",
    "revised results",
    "revised findings",
    "revised framing",
    "earlier draft",
    "current word manuscript",
    "raw archive",
    "project repository",
    "not currently recovered",
)


def update_tables(document: Document) -> None:
    table0 = document.tables[0]
    table0_rows = [
        ["Fit index", "chi-square (DoF)", "CFI", "TLI", "RMSEA", "AIC", "BIC"],
        ["One-factor capacity model", "741.72 (101)", "0.854", "0.818", "0.105", "67.41", "219.69"],
        ["Two-factor capacity model", "469.82 (98)", "0.915", "0.891", "0.081", "74.36", "239.69"],
    ]
    for r_idx, row in enumerate(table0_rows):
        for c_idx, value in enumerate(row):
            update_cell(table0.cell(r_idx, c_idx), value)

    table1 = document.tables[1]
    table1_rows = [
        ["Recovery outcome", "Predictor", "Standardized beta", "95% CI / p-value"],
        ["Recovery Timeliness", "Administrative Resources", "0.080", "[-0.004, 0.095]; p=0.072"],
        ["Recovery Timeliness", "Administrative Burden Capacity", "0.266", "[0.151, 0.352]; p<0.001"],
        ["Recovery Timeliness", "State or local government", "-0.106", "[-0.520, 0.000]; p=0.050"],
        ["Recovery Timeliness", "SVI Theme 1: Socioeconomic", "-0.194", "[-0.177, -0.035]; p=0.004"],
        ["Recovery Timeliness", "SVI Theme 2: Household composition & disability", "-0.159", "[-0.146, -0.027]; p=0.004"],
        ["Recovery Performance", "Administrative Resources", "-0.170", "[-0.212, -0.066]; p<0.001"],
        ["Recovery Performance", "Administrative Burden Capacity", "0.077", "[-0.015, 0.222]; p=0.086"],
        ["Recovery Performance", "State or local government", "-0.038", "[-0.526, 0.259]; p=0.505"],
        ["Recovery Performance", "SVI Theme 3: Minority status & language", "-0.226", "[-0.258, -0.096]; p<0.001"],
    ]
    for r_idx, row in enumerate(table1_rows):
        for c_idx, value in enumerate(row):
            update_cell(table1.cell(r_idx, c_idx), value)


def remove_table(table) -> None:
    """Remove a table from the DOCX body."""
    tbl = table._element
    tbl.getparent().remove(tbl)


def ensure_appendix_support_outputs() -> None:
    """Ensure the appendix support CSVs exist by rerunning recovered analysis if needed."""
    required = [
        KAIFA_DIAGNOSTICS_DIR / "recovered_notebook_measurement_appendix_table.csv",
        KAIFA_DIAGNOSTICS_DIR / "recovered_notebook_full_structural_reporting.csv",
        KAIFA_DIAGNOSTICS_DIR / "recovered_notebook_fit_verification_summary.csv",
        KAIFA_DIAGNOSTICS_DIR / "recovered_notebook_maturity_summary.csv",
        KAIFA_DIAGNOSTICS_DIR / "recovered_notebook_data_flow_summary.csv",
        KAIFA_DIAGNOSTICS_DIR / "recovered_notebook_sensitivity_summary.csv",
        KAIFA_DIAGNOSTICS_DIR / "recovered_notebook_ratio_artifact_summary.csv",
        KAIFA_DIAGNOSTICS_DIR / "recovered_notebook_official_geography_summary.csv",
        KAIFA_DIAGNOSTICS_DIR / "recovered_notebook_proxy_validation_summary.csv",
        KAIFA_DIAGNOSTICS_DIR / "recovered_notebook_subset_169_forensics_summary.csv",
    ]
    if all(path.exists() for path in required):
        return

    if str(SRC_ROOT) not in sys.path:
        sys.path.insert(0, str(SRC_ROOT))
    from capacity_sem.models.kaifa_recovered_analysis import run_recovered_analysis

    run_recovered_analysis(output_dir=KAIFA_DIAGNOSTICS_DIR, write_outputs=True)


def load_appendix_support_tables() -> dict[str, pd.DataFrame]:
    """Load appendix support tables written by the recovered Kaifa analysis stage."""
    ensure_appendix_support_outputs()
    return {
        "measurement": pd.read_csv(
            KAIFA_DIAGNOSTICS_DIR / "recovered_notebook_measurement_appendix_table.csv"
        ),
        "full_structural": pd.read_csv(
            KAIFA_DIAGNOSTICS_DIR / "recovered_notebook_full_structural_reporting.csv"
        ),
        "fit_verification": pd.read_csv(
            KAIFA_DIAGNOSTICS_DIR / "recovered_notebook_fit_verification_summary.csv"
        ),
        "maturity": pd.read_csv(
            KAIFA_DIAGNOSTICS_DIR / "recovered_notebook_maturity_summary.csv"
        ),
        "data_flow": pd.read_csv(
            KAIFA_DIAGNOSTICS_DIR / "recovered_notebook_data_flow_summary.csv"
        ),
        "geography": pd.read_csv(
            KAIFA_DIAGNOSTICS_DIR / "recovered_notebook_official_geography_summary.csv"
        ),
        "sensitivity": pd.read_csv(
            KAIFA_DIAGNOSTICS_DIR / "recovered_notebook_sensitivity_summary.csv"
        ),
        "ratio_artifacts": pd.read_csv(
            KAIFA_DIAGNOSTICS_DIR / "recovered_notebook_ratio_artifact_summary.csv"
        ),
        "proxy_validation": pd.read_csv(
            KAIFA_DIAGNOSTICS_DIR / "recovered_notebook_proxy_validation_summary.csv"
        ),
        "subset_forensics": pd.read_csv(
            KAIFA_DIAGNOSTICS_DIR / "recovered_notebook_subset_169_forensics_summary.csv"
        ),
    }


def format_cell_value(value) -> str:
    """Format appendix-table values compactly for DOCX output."""
    if pd.isna(value):
        return "-"
    if isinstance(value, (int,)):
        return str(value)
    if isinstance(value, float):
        if abs(value - round(value)) < 1e-9:
            return str(int(round(value)))
        return f"{value:.3f}"
    return str(value)


def append_dataframe_table(document: Document, caption: str, df: pd.DataFrame) -> None:
    """Append a caption paragraph and formatted table to the end of the DOCX."""
    document.add_paragraph("")
    document.add_paragraph(caption)
    table = document.add_table(rows=len(df) + 1, cols=len(df.columns))
    if document.tables:
        table.style = document.tables[0].style

    headers = [str(col) for col in df.columns]
    for idx, header in enumerate(headers):
        update_cell(table.cell(0, idx), header)

    for r_idx, (_, row) in enumerate(df.iterrows(), start=1):
        for c_idx, value in enumerate(row.tolist()):
            update_cell(table.cell(r_idx, c_idx), format_cell_value(value))


def append_supporting_appendix_tables(document: Document) -> None:
    """Replace the weak risk-ranking table with appendix support tables."""
    if len(document.tables) >= 3:
        remove_table(document.tables[2])

    support = load_appendix_support_tables()

    measurement_df = support["measurement"][
        [
            "latent_construct",
            "indicator",
            "std_loading",
            "std_error",
            "std_residual_var",
            "note",
        ]
    ].rename(
        columns={
            "latent_construct": "Latent construct",
            "indicator": "Indicator",
            "std_loading": "Std. loading",
            "std_error": "SE",
            "std_residual_var": "Std. residual var",
            "note": "Note",
        }
    )
    full_structural_df = support["full_structural"].rename(
        columns={
            "section": "Section",
            "lval": "Left side",
            "op": "Op",
            "rval": "Right side",
            "Estimate": "Estimate",
            "Est. Std": "Std. estimate",
            "Std. Err": "SE",
            "p-value": "p-value",
            "note": "Note",
        }
    )
    fit_verification_df = support["fit_verification"].rename(
        columns={
            "Statistic": "Fit statistic",
            "OneFactor": "Imported one-factor",
            "OneFactor_rerun_reference_ready": "Rerun one-factor",
            "Delta_OneFactor": "Delta one-factor",
            "TwoFactor": "Imported two-factor",
            "TwoFactor_rerun_reference_ready": "Rerun two-factor",
            "Delta_TwoFactor": "Delta two-factor",
            "note": "Note",
        }
    )
    maturity_df = support["maturity"].rename(
        columns={
            "maturity_band": "Portfolio maturity proxy",
            "jurisdictions": "N",
            "state_cases": "State",
            "local_cases": "Local",
            "median_programs": "Median programs",
            "median_disasters": "Median disasters",
            "median_duration_months": "Median duration (months)",
        }
    )
    sensitivity_df = support["sensitivity"][
        [
            "specification",
            "CFI",
            "RMSEA",
            "burden_to_timeliness_beta",
            "burden_to_performance_beta",
            "note",
        ]
    ].rename(
        columns={
            "specification": "Specification",
            "CFI": "CFI",
            "RMSEA": "RMSEA",
            "burden_to_timeliness_beta": "Burden -> timeliness beta",
            "burden_to_performance_beta": "Burden -> performance beta",
            "note": "Note",
        }
    )
    data_flow_df = support["data_flow"].rename(
        columns={
            "stage": "Stage",
            "count": "Count",
            "note": "Note",
        }
    )
    geography_df = support["geography"].rename(
        columns={
            "mapping_step": "Mapping step",
            "description": "Description",
            "count_or_implication": "Count / implication",
        }
    )
    proxy_validation_df = support["proxy_validation"].rename(
        columns={
            "check": "Check",
            "sample": "Sample",
            "estimate": "Estimate",
            "note": "Note",
        }
    )
    ratio_artifact_df = support["ratio_artifacts"].rename(
        columns={
            "element": "Element",
            "value": "Value",
            "note": "Note",
        }
    )
    subset_forensics_df = support["subset_forensics"].rename(
        columns={
            "finding": "Finding",
            "value": "Value",
            "note": "Note",
        }
    )

    document.add_paragraph("")
    document.add_paragraph("Appendix Tables")
    append_dataframe_table(
        document,
        "Appendix Table A1. Two-factor SEM measurement diagnostics.",
        measurement_df,
    )
    append_dataframe_table(
        document,
        "Appendix Table A2. Complete structural and latent-covariance reporting for the two-factor SEM.",
        full_structural_df,
    )
    append_dataframe_table(
        document,
        "Appendix Table A3. Imported-versus-rerun fit verification for the archived SEM models.",
        fit_verification_df,
    )
    append_dataframe_table(
        document,
        "Appendix Table A4. Recoverable data flow from standardized QPR rows to the archived 573-jurisdiction SEM sample.",
        data_flow_df,
    )
    append_dataframe_table(
        document,
        "Appendix Table A5. Portfolio maturity proxy summary for the 573-jurisdiction SEM sample.",
        maturity_df,
    )
    append_dataframe_table(
        document,
        "Appendix Table A6. Light-touch SEM sensitivity checks.",
        sensitivity_df,
    )
    append_dataframe_table(
        document,
        "Appendix Table A7. Ratio-artifact frequency and capped-ratio sensitivity summary.",
        ratio_artifact_df,
    )
    append_dataframe_table(
        document,
        "Appendix Table A8. Official 2023 Census geography crosswalk audit for SEM-facing units.",
        geography_df,
    )
    append_dataframe_table(
        document,
        "Appendix Table A9. Proxy-validation and coupling diagnostics for QCEW-based resource and workload measures.",
        proxy_validation_df,
    )
    append_dataframe_table(
        document,
        "Appendix Table A10. Forensic summary of the smaller cleaned N = 169 sensitivity sample.",
        subset_forensics_df,
    )


def validate_standalone_language() -> None:
    """Reject manuscript prose that reads like a revision memo."""
    lowered = "\n".join(PARAGRAPH_REPLACEMENTS.values()).lower()
    found = [phrase for phrase in FORBIDDEN_STANDALONE_PHRASES if phrase in lowered]
    if found:
        joined = ", ".join(sorted(found))
        raise ValueError(
            "Standalone manuscript validation failed; forbidden phrases found: "
            f"{joined}"
        )


def revise_document(document: Document) -> None:
    validate_standalone_language()
    for idx, text in PARAGRAPH_REPLACEMENTS.items():
        replace_paragraph_text(document.paragraphs[idx], text)
    for idx in [3, 7, 8]:
        normalize_empty_paragraph(document.paragraphs[idx])
    update_tables(document)
    append_supporting_appendix_tables(document)


def main() -> None:
    parser = argparse.ArgumentParser(description="Create a full Kaifa SEM manuscript DOCX.")
    parser.add_argument("--source", required=True, help="Path to the source DOCX manuscript.")
    parser.add_argument("--output", required=True, help="Path to write the manuscript DOCX.")
    args = parser.parse_args()

    source = Path(args.source)
    output = Path(args.output)

    document = Document(source)
    revise_document(document)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    print(f"Wrote manuscript to {output}")


if __name__ == "__main__":
    main()
