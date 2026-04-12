#!/usr/bin/env python3
"""
Post-process a rendered DOCX so it follows the formatting conventions of the
provided Kaifa manuscript source DOCX.

This is intentionally a formatting bridge, not a content source of truth.
Quarto remains the content source; the supplied Word manuscript remains the
formatting reference.
"""

from __future__ import annotations

import argparse
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt


def set_style_font(style, *, name="Aptos", size=12, bold=False, italic=False):
    """Apply basic font settings to a paragraph style."""
    font = style.font
    font.name = name
    font.size = Pt(size)
    font.bold = bold
    font.italic = italic


def remove_paragraph(paragraph):
    """Remove a paragraph from a document."""
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def insert_paragraph_after(paragraph, text="", style=None):
    """Insert a new paragraph after an existing paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_para._p = new_p
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def clear_paragraph(paragraph):
    """Remove all runs from a paragraph."""
    p = paragraph._element
    for child in list(p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            p.remove(child)


def ensure_style(document: Document, style_name: str):
    """Return a paragraph style by name if it exists."""
    try:
        return document.styles[style_name]
    except KeyError:
        return None


def restyle_document(document: Document):
    """Restyle the rendered document to approximate the supplied Word file."""
    normal = document.styles["Normal"]
    set_style_font(normal, name="Aptos", size=12)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name in ["Body Text", "First Paragraph", "Bibliography"]:
        style = ensure_style(document, style_name)
        if style is not None:
            style.base_style = normal
            set_style_font(style, name="Aptos", size=12)
            style.paragraph_format.space_after = Pt(8)
            style.paragraph_format.line_spacing = 1.15

    compact = ensure_style(document, "Compact")
    if compact is not None:
        compact.base_style = normal
        set_style_font(compact, name="Aptos", size=12)
        compact.paragraph_format.left_indent = Inches(0.5)
        compact.paragraph_format.space_after = Pt(0)
        compact.paragraph_format.line_spacing = 1.15

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = ensure_style(document, style_name)
        if style is not None:
            style.base_style = normal
            set_style_font(style, name="Aptos", size=12, bold=True)
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.line_spacing = 1.15

    title = ensure_style(document, "Title")
    if title is not None:
        title.base_style = normal
        set_style_font(title, name="Aptos", size=14, bold=True)
        title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(6)

    author = ensure_style(document, "Author")
    if author is not None:
        author.base_style = normal
        set_style_font(author, name="Aptos", size=12)
        author.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author.paragraph_format.space_after = Pt(0)

    for style_name in ["Image Caption", "Table Caption"]:
        style = ensure_style(document, style_name)
        if style is not None:
            style.base_style = normal
            set_style_font(style, name="Aptos", size=10)
            style.paragraph_format.space_before = Pt(4)
            style.paragraph_format.space_after = Pt(8)


def update_front_matter(document: Document, source_document: Document):
    """Make the rendered front matter follow the supplied Word manuscript."""
    paragraphs = document.paragraphs
    if len(paragraphs) < 6:
        return

    source_title = source_document.paragraphs[0].text.strip()
    source_authors = source_document.paragraphs[1].text.strip()
    source_affiliation = source_document.paragraphs[2].text.strip()
    source_keywords = source_document.paragraphs[6].text.strip()

    title_para = paragraphs[0]
    title_para.style = document.styles["Title"]
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    author_para = paragraphs[1]
    clear_paragraph(author_para)
    author_para.style = document.styles["Author"]
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_para.add_run(source_authors)

    # Remove extra per-author lines from pandoc output.
    remove_paragraph(document.paragraphs[2])
    remove_paragraph(document.paragraphs[2])

    author_para = document.paragraphs[1]
    affiliation_para = insert_paragraph_after(author_para, source_affiliation, style=document.styles["Normal"])
    affiliation_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Replace abstract heading with Word-manuscript formatting.
    abstract_heading = None
    for para in document.paragraphs:
        if para.text.strip() == "Abstract":
            abstract_heading = para
            break
    if abstract_heading is not None:
        clear_paragraph(abstract_heading)
        abstract_heading.style = document.styles["Normal"]
        run = abstract_heading.add_run("Abstract:")
        run.bold = True

    intro_para = None
    intro_idx = None
    for idx, para in enumerate(document.paragraphs):
        if para.text.strip() == "1. Introduction":
            intro_para = para
            intro_idx = idx
            break
    if intro_para is not None and intro_idx is not None:
        # Only insert keywords if not already present.
        if "Keywords:" not in "".join(p.text for p in document.paragraphs[:10]):
            abstract_body = document.paragraphs[intro_idx - 1]
            keywords_para = insert_paragraph_after(
                abstract_body,
                "",
                style=document.styles["Normal"],
            )
            first, rest = source_keywords.split(":", 1)
            key_run = keywords_para.add_run(first + ":")
            key_run.bold = True
            keywords_para.add_run(rest)

    # Match source convention for these heads.
    replacement_heads = {
        "9. Data Availability Statement": "Data Availability Statement",
        "10. References": "References",
    }
    for para in document.paragraphs:
        text = para.text.strip()
        if text in replacement_heads:
            clear_paragraph(para)
            para.style = document.styles["Normal"]
            run = para.add_run(replacement_heads[text])
            run.bold = True

    # Apply the source title if the rendered title was left generic.
    if title_para.text.strip() == "Title":
        clear_paragraph(title_para)
        run = title_para.add_run(source_title)
        run.bold = True


def main():
    parser = argparse.ArgumentParser(description="Post-process a rendered DOCX using the supplied Word manuscript format.")
    parser.add_argument("--source", required=True, help="Source Word manuscript used as formatting reference.")
    parser.add_argument("--input", required=True, help="Rendered input DOCX.")
    parser.add_argument("--output", required=True, help="Output DOCX path.")
    args = parser.parse_args()

    source_path = Path(args.source)
    input_path = Path(args.input)
    output_path = Path(args.output)

    source_doc = Document(source_path)
    rendered_doc = Document(input_path)

    restyle_document(rendered_doc)
    update_front_matter(rendered_doc, source_doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    rendered_doc.save(output_path)


if __name__ == "__main__":
    main()
